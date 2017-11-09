from docx.shared import Cm

def set_column_width(column, width):
  column.width = width
  for cell in column.cells:
    cell.width = width

def find_marker_table(doc,identifier):
  for t in doc.tables:
    top_left = t.rows[0].cells[0]
    if top_left.text == identifier:
      return t
  return None

def append_table(doc,dataframe):
  t = doc.add_table(1,1)
  populate_table(t,dataframe)

def replace_table(doc,identifier,dataframe):
  t = find_marker_table(doc,identifier)
  populate_table(t,dataframe)

def populate_table(t,dataframe):
  import pandas as pd
  n_index_cols = 1
  multi_index=False
  if isinstance(dataframe.index,pd.MultiIndex):
    multi_index=True
    n_index_cols = len(dataframe.index.levels)

  n_rows = len(dataframe) + 1
  n_cols = len(dataframe.columns) + n_index_cols

  while len(t.rows)<n_rows:
    t.add_row()
  while len(t.columns) < n_cols:
    t.add_column(Cm(2))

  # 
  for col in t.columns:
    set_column_width(col,Cm(2))

  header_row = t.rows[0]
  header_row.cells[0].text = ''
  for i,cell in enumerate(header_row.cells[n_index_cols:]):
    cell.text = dataframe.columns[i]
  t_columns = list(t.columns)
  index_cols = t_columns[0:n_index_cols]
  for i,index_col in enumerate(index_cols):
    for j,ix in enumerate(dataframe.index):
      cell = index_col.cells[j+1]
      if multi_index:
        here = ix[i]
        if here == dataframe.index[j-1][i]:
          cell.merge(index_col.cells[j])
          continue

        #if j > 0 and dataframe.index[j-1][i]
        cell.text = str(here)
        cell.paragraphs[0].runs[0].bold=True
      else:
        cell.text = str(ix)
        cell.paragraphs[0].runs[0].bold=True

  data_cols = t_columns[n_index_cols:]
  for df_col,t_col in zip(dataframe.columns,data_cols):
    series = dataframe[df_col]
#        print(i,series.iloc[0])
    for j,row in enumerate(dataframe.index):
#            print(i,j,series.iloc)
      cell = t_col.cells[j+1]
      cell.text = str(series.iloc[int(j)])

def insert_figure(doc,fn,marker_text,width=None,height=None,width_cm=None,height_cm=None):
  if width_cm:
    width = Cm(width_cm)
  
  if height_cm:
    height = Cm(height_cm)

  for para in doc.paragraphs:
    if para.text == marker_text:
      para.text = ''
      para.runs[0].add_picture(fn,width,height)